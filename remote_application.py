# -*- coding: utf-8 -*-
import socket
import threading
import time
import sys
import pymysql
import os
import numpy
import numpy as np
import PySide2
from PySide2.QtCore import QTimer
from PySide2.QtCore import QUrl
import pyqtgraph
import sys
import random
import pymysql
import pickle
import scipy.signal
import Ui_remote_main_ui
import Ui_remote_all_ui
import Ui_remote_applicantion_set_ui
import Ui_remote_sign_ui
import Ui_remote_vedio_player_ui
import docx
from queue import Queue,LifoQueue,PriorityQueue
#from PySide2.QtWidgets import QWidget, QPushButton, QApplication, QMessageBox, QGridLayout


import matplotlib
matplotlib.use('Qt5Agg')
from PySide2.QtWidgets import QMainWindow, QVBoxLayout, QWidget, QApplication
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg, NavigationToolbar2QT as NavigationToolbar
from matplotlib.figure import Figure
import matplotlib.pyplot as plt
import json
import pymysql

from PySide2 import QtWidgets, QtCore, QtGui, QtWebEngineWidgets
from PySide2.QtWebEngine import QtWebEngine
from PySide2.QtWidgets import QApplication
from PySide2.QtWebEngineWidgets import *

from PySide2.QtMultimedia import QMediaPlayer, QMediaPlaylist
from PySide2.QtMultimediaWidgets import QVideoWidget
# 将全局的字体设置为黑体
plt.rcParams['font.family'] = 'SimHei'
g_MaxBytes=1024*1024
QtCore.QCoreApplication.setAttribute(QtCore.Qt.AA_EnableHighDpiScaling)
QtCore.QCoreApplication.setAttribute(QtCore.Qt.AA_ShareOpenGLContexts)
QtWebEngine.initialize()
'''
/*
 * ***************************************************************************************
 * @brief 钢丝绳远程智能监控系统远程端程序，需配合现场端软件与服务端软件使用
 * @file remote_application.py
 * @author 张宏伟
 * @E-mail:bazhangya@hotmail.com
 * @QQ:515615456
 * @version V1.0
 * @data 2021-11-16
 * **************************************
 * @attention
 * 
 * 运行平台：windows平台，实测win10、win11可行
 * 运行环境：python3 pyside2
 * 运行前需要搭建运行环境，安装python3 pyside2和需要import的各种包，一般使用pycharm或者vscode，使用一个就行
 * @log
 * 2021-11-16开始注释并重构
 * 2021-11-17继续注释并重构
 * 2021-11-18重构、在线地点汇总逻辑
 * 2021-11-20完善在线地点显示与点击
 * ****************************************************************************************
**/
'''
'''
/*
 * @brief 主类，是自定义的，1年前写的耦合度很高，main函数只实例化这一个对象就能实现所有功能
 * @param 无
 * @retval 无
*/
'''
class main_application_obj():
    '''
    /*
    * @brief 类的构造函数，实例化对象时首先调用
    * @param self
    * @retval 无
    */
    '''
    def __init__(self):
        self.file_folder_init()
        self.variable_init()
        self.user_widget_init()
        self.signal_slot_init()
        self.graph_init()
        self.getNavTable()    
    '''
    /*
    * @brief 初始化存储数据的文件夹，构造函数中调用了一次
    * @param self
    * @retval 无
    */
    '''
    def file_folder_init(self):
        path = "C:\\detec_data"
        if not os.path.exists(path):
            print("文件夹不存在")
            os.makedirs(path)
        else:
            print("文件夹已存在")
    '''
    /*
    * @brief 子对象的初始化，诸如登录页面、报表页面、设置页面，生命周期贯穿始终，构造函数调用一次
    * @param self
    * @retval 无
    */
    '''
    def user_widget_init(self):
        ######窗口初始化
        self.main_window = PySide2.QtWidgets.QMainWindow()
        self.all_window = PySide2.QtWidgets.QMainWindow()
        self.set_window = PySide2.QtWidgets.QMainWindow()
        self.sign_widget = PySide2.QtWidgets.QWidget()
        self.video_widget = PySide2.QtWidgets.QWidget()
        ######ui对象
        self.main_ui = Ui_remote_main_ui.Ui_MainWindow()
        self.main_ui.setupUi(self.main_window)
        self.all_ui = Ui_remote_all_ui.Ui_MainWindow()
        self.all_ui.setupUi(self.all_window)
        self.set_ui = Ui_remote_applicantion_set_ui.Ui_MainWindow()
        self.set_ui.setupUi(self.set_window)
        self.sign_ui = Ui_remote_sign_ui.Ui_Form()
        self.sign_ui.setupUi(self.sign_widget)
        self.video_ui = Ui_remote_vedio_player_ui.Ui_Form()
        self.video_ui.setupUi(self.video_widget)
        self.sign_widget.showMaximized()
        self.sign_widget.show()
        
        #self.sign_widget.setFixedSize(self.sign_widget.width(), self.sign_widget.height())
        #self.all_window.setFixedSize(self.all_window.width(), self.all_window.height())
        #self.main_window.setFixedSize(self.main_window.width(), self.main_window.height())
        #self.all_window.show()
    '''
    /*
    * @brief 信号与槽的连接，构造函数调用一次
    * @param self
    * @retval 无
    */
    '''
    def signal_slot_init(self):
        self.main_ui.bn_start.clicked.connect(self.bn_slot_start)
        self.main_ui.bn_stop.clicked.connect(self.bn_slot_stop)
        self.main_ui.bn_report.clicked.connect(self.bn_slot_report)
        self.main_ui.bn_back.clicked.connect(self.bn_slot_back)
        self.main_ui.bn_set.clicked.connect(self.bn_slot_set)
        self.main_ui.bn_quit.clicked.connect(self.bn_slot_drop)
        #self.all_ui.pushButton_update.clicked.connect(self.bn_slot_update)
        #self.all_ui.pushButton_cancel.clicked.connect(self.bn_slot_sign_cancel) itemDoubleClicked
        self.all_ui.listWidget_summary.itemClicked.connect(self.listwidget_slot_signin_itemchoose)
        self.all_ui.listWidget_summary.itemDoubleClicked.connect(self.listwidget_slot_all_itemDoubleClicked)
        self.set_ui.pushButton_set_save.clicked.connect(self.bn_slot_set_save)
        self.set_ui.pushButton_set_cancel.clicked.connect(self.bn_slot_set_cancel)
        self.set_ui.pushButton_set_signin.clicked.connect(self.bn_slot_set_signin)
        self.set_ui.comboBox_set_servechose.currentIndexChanged.connect(self.cb_slot_set_servechoose)
        self.sign_ui.pushButton_sign.clicked.connect(self.bn_slot_sign_sign)
        self.sign_ui.pushButton_quit.clicked.connect(self.bn_slot_sign_quit)
        self.video_ui.pushButton_vedio_quit.clicked.connect(self.bn_slot_vedio_quit)
        self.video_ui.pushButton_vedio_report.clicked.connect(self.bn_slot_vedio_report)
        
    '''
    /*
    * @brief socket通信函数，连接采集卡，没用上，本地端的代码，只是没删除
    * @param self
    * @retval 无
    */
    '''
    def socket_init(self):######以太网通信对象
        self.tcp_flag = 1
        self.tcp_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        # 2. 链接服务器
        self.tcp_socket.connect(("192.168.1.10", 1000))
    '''
    /*
    * @brief socket通信函数，连接服务器
    * @param self
    * @retval 无
    */
    '''
    def internet_socket_init(self):######服务器连接
        self.internet_tcp_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        # 2. 链接服务器
        self.internet_tcp_socket.connect(("118.178.187.119", 8889))
        self.internet_tcp_send_command = "get," + self.baseName
        self.internet_tcp_socket.sendall(self.internet_tcp_send_command.encode())
    '''
    /*
    * @brief 连接服务器数据库的函数，没用上
    * @param self
    * @retval 无
    */
    '''
    def userpysql_init(self):#######数据库对象
        self.mysql_db = pymysql.connect('118.178.187.119',port=3306,user = self.mysql_id,passwd = self.mysql_password,db = 'Remote_Detection')
        self.mysql_cursor = self.mysql_db.cursor()
        self.mysql_sql ="insert into wh_detec(channel0,channel1, channel2, channel3, channel4,channel5,channel6, channel7, channel8, channel9, channel10, channel11, channel12,channel13,channel14, channel15) values(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
    '''
    /*
    * @brief 绘图pyqtgraph与pyside2的连接
    * @param self
    * @retval 无
    */
    '''
    def graph_init(self):
        ######绘图对象
        pyqtgraph.setConfigOption('background', (6, 32, 69,50))
        pyqtgraph.setConfigOption('foreground', 'w')
        self.graph_layout = pyqtgraph.GraphicsLayoutWidget()
        self.coil_plot = self.graph_layout.addPlot(colspan=2)
        self.coil_plot.showGrid(x=True, y=True)
        self.coil_plot.setLabel(axis="left", text="coil voltage")
        self.coil_plot.setLabel(axis="bottom", text="time")
        #self.data_plot.setTitle("钢丝绳损伤检测")
        self.coil_plot.setRange(yRange=[self.set_graph_up_ruler_limit, self.set_graph_down_ruler_limit])
        self.coil_plot.addLegend()
        self.coil_curve1 = self.coil_plot.plot(pen="w", name="y1")
        self.coil_curve2 = self.coil_plot.plot(pen='r', name="y2")
        self.main_ui.third_layer2.addWidget(self.graph_layout)
        self.graph_layout.nextRow()####下一行
        self.hal_plot = self.graph_layout.addPlot(colspan=2)
        self.hal_plot.showGrid(x=True, y=True)
        self.hal_plot.setLabel(axis="left", text="hal voltage")
        self.hal_plot.setLabel(axis="bottom", text="time")
        # self.data_plot.setTitle("钢丝绳损伤检测")
        self.hal_plot.setRange(yRange=[self.set_graph_up_ruler_limit, self.set_graph_down_ruler_limit])
        self.hal_plot.addLegend()
        self.hal_curve1 = self.hal_plot.plot(pen="w", name="y1")
        self.hal_curve2 = self.hal_plot.plot(pen='r', name="y2")
        self.coil_curve1.setData(self.graph_coil_plot_data1)
        self.coil_curve1.setPos(self.graph_coil_plot_pos1, 0)
        self.coil_curve2.setData(self.graph_coil_plot_data2)
        self.coil_curve2.setPos(self.graph_coil_plot_pos2, 0)
        self.hal_curve1.setData(self.graph_hal_plot_data1)
        self.hal_curve1.setPos(self.graph_hal_plot_pos1, 0)
        self.hal_curve2.setData(self.graph_hal_plot_data2)
        self.hal_curve2.setPos(self.graph_hal_plot_pos2, 0)

        #播放视频
        self.videoWidget = QVideoWidget()
        self.mediaPlayer = QMediaPlayer()
        #self.mediaPlaylist = QMediaPlaylist()
        # Add the video file path   
        #self.mediaPlaylist.addMedia(QUrl.fromLocalFile(os.path.abspath("D:/系统默认/IDM下载/MediaPlayer-master/sample_data/sampleVideo.mp4")))
        # Set the video to played in a loop once it ends.
        #self.mediaPlaylist.setPlaybackMode(QMediaPlaylist.CurrentItemInLoop)
        # Set the QMediaPlaylist for the QMediaPlayer.
        #self.mediaPlayer.setPlaylist(self.mediaPlaylist)
        # Add the QVideoWidget in the GridLayout.
        self.video_ui.verticalLayout.addWidget(self.videoWidget)
        #self.playerLayout.addWidget(self.videoWidget)
        # Set the video output from the QMediaPlayer to the QVideoWidget.
        self.mediaPlayer.setVideoOutput(self.videoWidget)
        # Set the QPushButtons to play, pause and stop the video in the QVideoWidget.
        #地图
        self.view = QWebEngineView()
        #self.all_ui.verticalLayout_map.addWidget(self.view)
        #self.view.load(QUrl.fromLocalFile(os.path.abspath("E:/myhtml/main.html")))
        #self.view.show()
        #主界面的图形更新，柱状图
        self.sc = MplCanvas(self,facecolor = '#0a2461')
        '''
        self.sc.fig.clf()
        self.sc.fig.patch.set_facecolor('#0a2461')
        self.sc.fig.patch.set_alpha(1)
        self.sc.axes = self.sc.fig.add_subplot(111)
        self.sc.axes.tick_params(axis='x', colors='w')    #setting up X-axis tick color to red
        self.sc.axes.tick_params(axis='y', colors='w')  #setting up Y-axis tick color to black
        self.sc.axes.spines['left'].set_color('w')        # setting up Y-axis tick color to red
        self.sc.axes.spines['bottom'].set_color('w')        # setting up Y-axis tick color to red
        self.sc.axes.spines['top'].set_color('w')        # setting up Y-axis tick color to red
        self.sc.axes.spines['right'].set_color('w')        # setting up Y-axis tick color to red
        x = [1,2,3,4,5]
        y = [15,2,5,4,8]
        labe = ['设备1','设备2','设备3','设备4','设备5']
        self.sc.axes.bar(x, y, alpha=1, width=0.3, color='#db4300', edgecolor='w',tick_label = labe)
        self.sc.axes.set(facecolor = "#0a2461")
        self.sc.axes.set(alpha = 0)
        '''
        self.all_ui.verticalLayout_graph.addWidget(self.sc)
        #主界面的图形更新，饼图
        self.graphCake = MplCanvas(self,facecolor = '#000928')
        '''
        self.graphCake.fig.clf()
        self.graphCake.fig.patch.set_facecolor('#000928')
        self.graphCake.fig.patch.set_alpha(1)
        self.graphCake.axes = self.graphCake.fig.add_subplot(111)
        # Pie chart, where the slices will be ordered and plotted counter-clockwise:
        labels = '在检', '未检', '离线'
        sizes = [1, 2, 2]
        explode = (0.1, 0, 0)  # only "explode" the 2nd slice (i.e. 'Hogs')
        _, _, autotexts = self.graphCake.axes.pie(sizes, explode=explode, labels=labels,autopct='%1.1f%%',shadow=True, startangle=90,textprops={'color':"w"})
        self.graphCake.axes.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
        '''
        self.all_ui.verticalLayout_cake.addWidget(self.graphCake)
        self.graph_update_visualization(self.g_damageStats)
    '''
    /*
    * @brief 一些变量、对象或者标志位的初始化
    * @param self
    * @retval 无
    */
    '''
    def variable_init(self):
        ######TCP通信多线程flag
        self.baseName = "null"
        self.thread_socket_flag = 1
        self.thread_defect_judge_flag = 1
        self.thread_socket_frequency_print = 0
        self.thread_tcp_commucate = tcp_commucate_qtThread()
        self.thread_defect_judge_obj = defect_judge_qtThread()
        self.thread_defect_buff_data = []
        self.thread_defect_judge_data = []
        self.thread_defect_file_save_data = []
        self.thread_defect_data_copy_flag = 0
        self.thread_defect_hurt_number = 0
        self.thread_defect_coder_location = 0
        self.set_up_damage_limit = 4000
        self.set_down_damage_limit = 0
        self.set_graph_up_ruler_limit = 4000
        self.set_graph_down_ruler_limit = 0
        self.set_defect_judge_interval = 3200
        self.tcp_socket = 0  #####
        self.internet_tcp_socket = 0
        self.internet_tcp_send_command = "get,"
        self.tcp_receive_data = [0 for x in range(25600)]
        self.tcp_queue_user1_data=Queue(maxsize=40)#先进先出队列
        self.mysql_db = 0
        self.mysql_cursor = 0
        self.mysql_sql = 0
        self.mysql_id = "Remote_Detection"
        self.mysql_password =  "Hfs6deE6KnYcyA6C"
        self.graph_coil_plot_data1 = [1900 for x in range(8000)]
        self.graph_coil_plot_pos1 = 0
        self.graph_coil_plot_data2 = [2100 for x in range(8000)]
        self.graph_coil_plot_pos2 = 0
        self.graph_hal_plot_data1 = [1900 for x in range(8000)]
        self.graph_hal_plot_pos1 = 0
        self.graph_hal_plot_data2 = [2100 for x in range(8000)]
        self.graph_hal_plot_pos2 = 0
        self.graph_coil_sub_data1 = 0
        self.graph_coil_sub_data2 = 0
        self.graph_hal_sub_data1 = 0
        self.graph_hal_sub_data2 = 0
        self.graph_sub_data_run_once_flag = 1
        self.graph_update_timer = pyqtgraph.QtCore.QTimer()
        self.QTimer_Graph_Update = QTimer()
        self.internet_flag = 0
        self.g_damageStats = {'设备1': ['设备1', '武汉', 15, 0], '设备2': ['设备2', '武汉', 20, 1], '设备3': ['设备3', '武汉', 5, 0], '设备4': ['设备4', '武汉', 5, 0]}
        self.json_command = {"mod":"remote","name":"all","base":"武汉","damageNum":0,"flag_cloud_uploading":0}
        self.sc = 0
        self.graphCake = 0
        self.flag_vedio_widget = 0
        self.startTime = time.time()
    '''
    /*
    * @brief 获取所有的监测点名字,
    * @param self
    * @retval 无
    */
    '''
    def getNavTable(self):
        try:
            self.internet_tcp_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            # 2. 链接服务器
            self.internet_tcp_socket.connect(("118.178.187.119", 8889))
            sendStr = json.dumps(self.json_command)
            self.internet_tcp_socket.sendall(sendStr.encode('utf-8'))
            navtable = self.internet_tcp_socket.recv(g_MaxBytes)
            self.internet_tcp_socket.close()
            self.g_damageStats = json.loads(navtable.decode('utf-8'))
            #print(self.g_damageStats)
            self.all_ui.listWidget_summary.clear()
            self.all_ui.listWidget_base.clear()
            self.all_ui.listWidget_num.clear()
            if self.g_damageStats == {}:
                print("当前没有在线监测设备")
                self.all_ui.listWidget_summary.addItem("无检测设备在线！")
            else:
                for key in self.g_damageStats:
                    self.all_ui.listWidget_summary.addItem(self.g_damageStats[key][0])
                    self.all_ui.listWidget_base.addItem(self.g_damageStats[key][1])
                    self.all_ui.listWidget_num.addItem(str(self.g_damageStats[key][2]))               
        except:
            #self.all_ui.listWidget_summary.addItem("网络链路出错")
            print("get navtable网络链路出错")
    ''' 
    /*
    * @brief 槽函数，点击开始按钮后自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_start(self):  #####开始按钮函数
        try:
            self.internet_tcp_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            # 2. 链接服务器
            self.internet_tcp_socket.connect(("118.178.187.119", 8889))
            self.internet_tcp_socket.sendall("get,navtable".encode())
            navtable = self.internet_tcp_socket.recv(1280)
            self.internet_tcp_socket.close()
            print(navtable.decode('utf-8'))
        except:
            self.bn_slot_drop()
            return
            print("网络有点问题")
        if navtable.decode('utf-8') == 'null':
            self.bn_slot_drop()
            return
            print("当前没有在线监测设备")
            self.all_ui.listWidget_summary.addItem("无检测设备在线！")
        else:
            baseNameList = navtable.decode('utf-8').split(",")
            print(baseNameList)
        if self.baseName not in baseNameList:
            self.bn_slot_drop()
            return
        self.main_ui.bn_start.setText("检测中")
        self.thread_defect_judge_flag = 1
        self.internet_flag = 1
        try:
            self.internet_socket_init()  ######服务器连接
            self.thread_tcp_commucate.start()
            self.thread_defect_judge_obj.start()
            self.graph_update_timer.timeout.connect(self.graph_update)
            self.graph_update_timer.start(8)
        except:
            self.main_ui.bn_start.setText("出错")
        print('检测')
    '''
    /*
    * @brief 测试函数
    * @param self
    * @retval 无
    */
    '''
    def test(self):
        print("ceshi1")
    '''
    /*
    * @brief 槽函数，点击停止按钮后自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_stop(self):  #####暂停按钮
        self.main_ui.bn_start.setText("检测")
        self.internet_flag = 0
        self.thread_socket_flag = 0
        self.thread_defect_judge_flag = 0
        self.graph_update_timer.stop()
        self.QTimer_Graph_Update.stop()
        print('暂停')
    '''
    /*
    * @brief 按钮槽函数，点击报表按钮后自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_report(self):  #####报表函数
        fileName_choose, filetype = PySide2.QtWidgets.QFileDialog.getOpenFileName(self.main_window, "选取文件", "C:\\detec_data","All Files (*);;Text Files (*.txt)")  # 设置文件扩展名过滤,用双分号间隔
        if fileName_choose == "":
            return
        pickle_file = open(fileName_choose, "rb")
        read_data = pickle.load(pickle_file)
        pickle_file.close()
        self.report_defect_docx_save(read_data)
    '''
    /*
    * @brief 按钮槽函数，点击回放按钮后自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_back(self):  #####回放按钮函数
        fileName_choose, filetype = PySide2.QtWidgets.QFileDialog.getOpenFileName(self.main_window, "选取文件","C:\\detec_data","All Files (*);;Text Files (*.txt)")  # 设置文件扩展名过滤,用双分号间隔
        if fileName_choose == "":
            return
        print("\n你选择的文件为:")
        print(fileName_choose)
        print("文件筛选器类型: ", filetype)
    '''
    /*
    * @brief 按钮槽函数，点击设置按钮后自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_set(self):  #####设置按钮函数
        self.set_window.show()
        self.main_window.close()
        self.set_ui.lineEdit_set_graph_y_up.setText(str(self.set_graph_up_ruler_limit))
        self.set_ui.lineEdit_set_graph_y_down.setText(str(self.set_graph_down_ruler_limit))
        self.set_ui.lineEdit_set_alarm_uplimit.setText(str(self.set_up_damage_limit))
        self.set_ui.lineEdit_set_alarm_down_limit.setText(str(self.set_down_damage_limit))
        print('设置')
    '''
    /*
    * @brief 按钮槽函数，点击退出按钮后自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_drop(self):  #####退出按钮函数
        self.thread_socket_flag = 0
        self.thread_defect_judge_flag = 0
        self.graph_update_timer.stop()
        self.all_window.show()
        self.main_window.close()
        print('退出')
    '''
    /*
    * @brief 按钮槽函数，点击登录按钮后自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_update(self):
        self.getNavTable()
        self.graph_update_visualization(self.g_damageStats)
    '''
    /*
    * @brief 按钮槽函数，点击取消按钮后自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_sign_cancel(self):
        self.all_window.close()
    '''
    /*
    * @brief 按钮槽函数，点击设置页面保存按钮后自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_set_save(self):
        self.set_graph_up_ruler_limit = int(self.set_ui.lineEdit_set_graph_y_up.text() )
        self.set_graph_down_ruler_limit = int(self.set_ui.lineEdit_set_graph_y_down.text() )
        self.set_up_damage_limit = int(self.set_ui.lineEdit_set_alarm_uplimit.text() )
        self.set_down_damage_limit = int(self.set_ui.lineEdit_set_alarm_down_limit.text() )
        self.coil_plot.setRange(yRange=[self.set_graph_up_ruler_limit, self.set_graph_down_ruler_limit])
        self.hal_plot.setRange(yRange=[self.set_graph_up_ruler_limit, self.set_graph_down_ruler_limit])
        self.set_window.close();self.main_window.show()
    '''
    /*
    * @brief 按钮槽函数，点击设置页面取消按钮后自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_set_cancel(self):
        self.set_window.close();self.main_window.show()
    '''
    /*
    * @brief 按钮槽函数，点击设置页面登录按钮后自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_set_signin(self):
        self.main_window.close()
        self.set_window.close()
        self.all_window.show()
    '''
    /*
    * @brief 按钮槽函数，点击登录按钮自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_sign_sign(self):
        time.sleep(0.3)
        my_ID = self.sign_ui.lineEdit_id.text()
        my_Passwd = self.sign_ui.lineEdit_passwd.text()
        if my_ID == 'admin' and my_Passwd == 'admin':
            self.all_window.showMaximized()
            self.sign_widget.close()
            self.all_window.show()
            self.QTimer_Graph_Update.timeout.connect(self.bn_slot_update)
            self.QTimer_Graph_Update.start(1000)
        else:
            PySide2.QtWidgets.QMessageBox.about(self.sign_widget, '温馨提示', '账号或密码错误，请检查后重试')

        
    '''
    /*
    * @brief 按钮槽函数，点击退出按钮自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_sign_quit(self):
        self.sign_widget.close()
    '''
    /*
    * @brief 按钮槽函数，点击按钮自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_vedio_quit(self):
        self.mediaPlayer.stop()
        self.video_widget.close()
        pass
    '''
    /*
    * @brief 按钮槽函数，点击按钮自动执行
    * @param self
    * @retval 无
    */
    '''
    def bn_slot_vedio_report(self):
        if self.flag_vedio_widget == 0:
            self.flag_vedio_widget = 1
            self.video_ui.pushButton_vedio_report.setText("暂停")
            self.mediaPlayer.play()
        else:
            self.flag_vedio_widget = 0
            self.video_ui.pushButton_vedio_report.setText("回放")
            self.mediaPlayer.pause()


    '''
    /*
    * @brief coxbox槽函数，点击设置页面选择服务器按钮后自动执行
    * @param self
    * @retval 无
    */
    '''
    def cb_slot_set_servechoose(self,i):
        self.internet_tcp_send_command = self.set_ui.comboBox_set_servechose.currentText()
        print(self.internet_tcp_send_command)
    '''
    /*
    * @brief 多线程，循环向服务端发送"get"获取数据
    * @param self
    * @retval 无
    */
    '''
    '''
    /*
    * @brief 槽函数，主页面，点击当前在线设备中的listwidgetitem自动执行
    * @param self item
    * @retval 无
    */
    '''
    def listwidget_slot_all_itemDoubleClicked(self,item):
        print(item.text())
        if item.text() == "网络链路出错":
            return
        if item.text() == "无检测设备在线！":
            return
        self.video_widget.showMaximized()
        self.video_widget.show()
        self.video_ui.pushButton_vedio_report.setEnabled(True)
        time.sleep(0.3)
        if item.text() == "喻远智能":
            self.mediaPlayer.setMedia(QUrl.fromLocalFile(os.path.abspath("E:/video/0525-4.mp4")))
            self.mediaPlayer.play()
            self.mediaPlayer.pause()
        elif item.text() == "实验室":
            self.mediaPlayer.setMedia(QUrl.fromLocalFile(os.path.abspath("E:/video/0525-3.mp4")))
            self.mediaPlayer.play()
            self.mediaPlayer.pause()
        elif item.text() == "手持一体":
            self.mediaPlayer.setMedia(QUrl.fromLocalFile(os.path.abspath("E:/video/0526-1.mp4")))
            self.video_ui.pushButton_vedio_report.setEnabled(False)
            self.mediaPlayer.play()
            self.mediaPlayer.pause()
        elif item.text() == "维尔卡":
            self.mediaPlayer.setMedia(QUrl.fromLocalFile(os.path.abspath("E:/video/0526-1.mp4")))
            self.video_ui.pushButton_vedio_report.setEnabled(False)
            self.mediaPlayer.play()
            self.mediaPlayer.pause()
        elif item.text() == "口孜东":
            self.mediaPlayer.setMedia(QUrl.fromLocalFile(os.path.abspath("E:/video/0526-2.mp4")))
            self.video_ui.pushButton_vedio_report.setEnabled(False)
            self.mediaPlayer.play()
            self.mediaPlayer.pause()
        elif item.text() == "伊化矿":
            self.mediaPlayer.setMedia(QUrl.fromLocalFile(os.path.abspath("E:/video/0525-5.mp4")))
            self.mediaPlayer.play()
            self.mediaPlayer.pause()
        #self.all_window.close()
    '''
    /*
    * @brief 槽函数，双击动作
    * @param self item
    * @retval 无
    */
    '''
    def listwidget_slot_signin_itemchoose(self,item):
        self.all_ui.verticalLayout_map.addWidget(self.view)
        if item.text() == "实验室":
            self.view.load(QUrl.fromLocalFile(os.path.abspath("E:/myhtml/shiyanshi.html")))
        elif item.text() == "喻远智能":
            self.view.load(QUrl.fromLocalFile(os.path.abspath("E:/myhtml/yuyuanzhineng.html")))
        elif item.text() == "维尔卡":
            self.view.load(QUrl.fromLocalFile(os.path.abspath("E:/myhtml/weierka.html")))
        elif item.text() == "手持一体":
            self.view.load(QUrl.fromLocalFile(os.path.abspath("E:/myhtml/yuyuanzhineng.html")))
        elif item.text() == "口孜东":
            self.view.load(QUrl.fromLocalFile(os.path.abspath("E:/myhtml/kouzidong.html")))
        elif item.text() == "伊化矿":
            self.view.load(QUrl.fromLocalFile(os.path.abspath("E:/myhtml/yihua.html")))
        self.view.show()
        print("signal clicked")
        pass
        
    def thread_tcp_commucate_get_func(self):
        while self.internet_flag:
            self.thread_socket_frequency_print += 1
            #print(self.thread_socket_frequency_print)
            send_data = self.internet_tcp_send_command #get,baseName
            self.internet_tcp_socket.sendall(send_data.encode())
            rece_data = list(self.internet_tcp_socket.recv(1280) )
            if len(rece_data) == 1280:
                trans_data = self.tcp_data_trs_together_func(rece_data)
                if self.tcp_queue_user1_data.full():
                    _ = self.tcp_queue_user1_data.get()
                else:
                    self.tcp_queue_user1_data.put(trans_data)
                self.thread_defect_buff_data.extend(trans_data)
                #self.graph_update_plot_func(trans_data)
            else:
                print("等待%d",self.thread_socket_frequency_print)
                PySide2.QtCore.QThread.msleep(6)
            PySide2.QtCore.QThread.msleep(6)
        self.internet_tcp_socket.close()
    '''
    /*
    * @brief 多线程，循环判断损伤
    * @param self
    * @retval 无
    */
    '''
    def thread_defect_judge_func(self):
        while self.thread_defect_judge_flag:
            if self.thread_defect_data_copy_flag == 1:
                #print(len(self.thread_defect_judge_data))
                for each in self.thread_defect_judge_data[0::4]:
                    if each > self.set_up_damage_limit or each < self.set_down_damage_limit:
                        self.thread_defect_hurt_number += 1
                        break
                for each in self.thread_defect_judge_data[1::4]:
                    if each > self.set_up_damage_limit or each < self.set_down_damage_limit:
                        self.thread_defect_hurt_number += 1
                        break
                self.main_ui.label_damage.setText(str(self.thread_defect_hurt_number))
                self.thread_defect_file_save_data.extend(self.thread_defect_judge_data)
                self.thread_defect_data_copy_flag = 0
                if self.graph_sub_data_run_once_flag == 1:#######将基线拉到2000
                    self.graph_sub_data_run_once_flag = 0
                    self.graph_coil_sub_data1 = int(numpy.mean(self.thread_defect_judge_data[0::4]) - 2000 )
                    self.graph_coil_sub_data2 = int(numpy.mean(self.thread_defect_judge_data[1::4]) - 2000 )
            else:
                PySide2.QtCore.QThread.msleep(2)
        self.thread_file_save_func(self.thread_defect_file_save_data)
    '''
    /*
    * @brief 多线程，循环判断损伤
    * @param self input_data
    * @retval 无
    */
    '''
    def graph_update_plot_func(self,input_data):#####用于绘制曲线
        self.graph_coil_plot_data1[:-40] = self.graph_coil_plot_data1[40:]  # shift data in the array one sample left
        self.graph_coil_plot_pos1 += 40
        self.graph_coil_plot_data1[-40:] = input_data[0::4]  # butter_lowpass_filter(input_data, 100, 1000, 5) input_data
        self.graph_coil_plot_data2[:-40] = self.graph_coil_plot_data2[40:]  # shift data in the array one sample left
        self.graph_coil_plot_pos2 += 40
        self.graph_coil_plot_data2[-40:] = input_data[1::4]  # butter_lowpass_filter(input_data, 100, 1000, 5) input_data
        self.graph_hal_plot_data1[:-40] = self.graph_hal_plot_data1[40:]  # shift data in the array one sample left
        self.graph_hal_plot_pos1 += 40
        self.graph_hal_plot_data1[-40:] = input_data[2::4]  # butter_lowpass_filter(input_data, 100, 1000, 5) input_data
        self.graph_hal_plot_data2[:-40] = self.graph_hal_plot_data2[40:]  # shift data in the array one sample left
        self.graph_hal_plot_pos2 += 40
        self.graph_hal_plot_data2[-40:] = input_data[3::4]  # butter_lowpass_filter(input_data, 100, 1000, 5) input_data
        self.coil_curve1.setData(self.graph_coil_plot_data1)
        self.coil_curve1.setPos(self.graph_coil_plot_pos1, 0)
        self.coil_curve2.setData(self.graph_coil_plot_data2)
        self.coil_curve2.setPos(self.graph_coil_plot_pos2, 0)
        self.hal_curve1.setData(self.graph_hal_plot_data1)
        self.hal_curve1.setPos(self.graph_hal_plot_pos1, 0)
        self.hal_curve2.setData(self.graph_hal_plot_data2)
        self.hal_curve2.setPos(self.graph_hal_plot_pos2, 0)
    '''
    /*
    * @brief 数据可视化的定时更新，定时器定时调用此函数
    * @param self 还有很多
    * @retval 无
    */
    '''
    def graph_update_visualization(self,g_damageStats):#####用于绘制曲线
         #主界面的图形更新，柱状图
        if len(g_damageStats) == 0:
            return
        self.sc.fig.clf()
        self.sc.fig.patch.set_facecolor('#0a2461')
        self.sc.fig.patch.set_alpha(1)
        self.sc.axes = self.sc.fig.add_subplot(111)
        self.sc.axes.tick_params(axis='x', colors='w')    #setting up X-axis tick color to red
        self.sc.axes.tick_params(axis='y', colors='w')  #setting up Y-axis tick color to black
        self.sc.axes.spines['left'].set_color('w')        # setting up Y-axis tick color to red
        self.sc.axes.spines['bottom'].set_color('w')        # setting up Y-axis tick color to red
        self.sc.axes.spines['top'].set_color('w')        # setting up Y-axis tick color to red
        self.sc.axes.spines['right'].set_color('w')        # setting up Y-axis tick color to red
        #x = [1,2,3,4,5]
        #y = [15,2,5,4,8]
        #labe = ['设备1','设备2','设备3','设备4','设备5']
        x=list(range(1,len(g_damageStats)+1))
        y=[]
        labe=[]
        onWorking = 0
        notWrking = 0
        offLine = 0
        damageNum = 0
        for key in g_damageStats:
            damageNum = damageNum + g_damageStats[key][2]
            y.append(g_damageStats[key][2])
            labe.append(key)
            if(g_damageStats[key][3] == 0):
                notWrking = notWrking+1
            else:
                onWorking = onWorking + 1
        #offLine = notWrking+onWorking
        offLine = 0
        print(x)
        print(y)
        print(labe)
        #x = [1,2,3,4,5]
        #y = [1,2,3,4,5]
        #labe = ["s","s","s","s","s"]
        
        self.sc.axes.bar(x, y, alpha=1, width=0.3, color='#db4300', edgecolor='w',tick_label = labe)
        self.sc.axes.set(facecolor = "#0a2461")
        self.sc.axes.set(alpha = 0)
        self.sc.figure.canvas.draw()#重绘，不然不更新图片
        self.all_ui.verticalLayout_graph.addWidget(self.sc)
        #主界面的图形更新，饼图
        self.graphCake.fig.clf()
        self.graphCake.fig.patch.set_facecolor('#000928')
        self.graphCake.fig.patch.set_alpha(1)
        self.graphCake.axes = self.graphCake.fig.add_subplot(111)
        # Pie chart, where the slices will be ordered and plotted counter-clockwise:
        labels = '在检', '未检', '离线'
        sizes = [onWorking,notWrking,offLine]
        explode = (0.1, 0, 0)  # only "explode" the 2nd slice (i.e. 'Hogs')
        _, _, autotexts = self.graphCake.axes.pie(sizes, explode=explode, labels=labels,autopct='%1.1f%%',shadow=True, startangle=90,textprops={'color':"w"})
        self.graphCake.axes.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
        self.graphCake.figure.canvas.draw()#重绘，不然不更新图片
        self.all_ui.verticalLayout_cake.addWidget(self.graphCake)

        self.all_ui.label_online.setText(str((onWorking+notWrking)))
        self.all_ui.label_offline.setText(str(offLine))
        runTime = 83.2+(time.time() - self.startTime)/3600
        self.all_ui.label_workingtime.setText(str(round(runTime, 4)))
        self.all_ui.label_damage.setText(str(damageNum))
        self.all_ui.label_onworking.setText(str(onWorking))
        self.all_ui.label_notworking.setText(str(notWrking))
        self.all_ui.label_offline.setText(str(offLine))
    '''
    /*
    * @brief 数据解包函数，高低八位整合
    * @param self source_data
    * @retval 无
    */
    '''
    def tcp_data_trs_func(self,source_data):   #'''高低八位数据整合函数
        return_data = []
        for each in range(640):
            return_data.append(source_data[2 * each] + source_data[2 * each + 1] * 256)
        return return_data
    '''
    /*
    * @brief 多线程，循环存储文件
    * @param self input_data
    * @retval 无
    */
    '''
    def thread_file_save_func(self,input_data):   #文件存储
        localtime = str(time.asctime(time.localtime(time.time())) )
        filename = localtime.replace(" ", "-")+'.pkl'
        filename = "C:\\detec_data\\"+filename.replace(":","_")
        pickle_file = open(filename, "wb")
        pickle.dump(input_data, pickle_file)
        pickle_file.close()
    '''
    /*
    * @brief 数据解包函数，高低八位整合，并且合并通道
    * @param self source_data
    * @retval 无
    */
    '''
    def tcp_data_trs_together_func(self,source_data):#高低八位数据整合函数,返回合合并通道数据
        pr_data = []
        return_data = []
        for i in range(640):
            pr_data.append(source_data[2 * i] + source_data[2 * i + 1] * 256)
        for each in range(40):
            return_data.extend([pr_data[each * 16 + 1] + pr_data[each * 16 + 2] + pr_data[each * 16 + 3] + pr_data[each * 16 + 4] + pr_data[each * 16 + 5] + pr_data[each * 16 + 6] - 10000,pr_data[each * 16 + 7] + pr_data[each * 16 + 8] + pr_data[each * 16 + 13] + pr_data[each * 16 + 14] + pr_data[each * 16 + 15] + pr_data[each * 16 + 0] - 10000, pr_data[each * 16 + 10], pr_data[each * 16 + 11]])
        return return_data
    '''
    /*
    * @brief 低通滤波函数
    * @param self cutoff fs order=5
    * @retval 无
    */
    '''
    def butter_lowpass(self,cutoff, fs, order=5):
        nyq = 0.5 * fs
        normal_cutoff = cutoff / nyq
        b, a = scipy.signal.butter(order, normal_cutoff, btype='low', analog=False)
        return b, a
    '''
    /*
    * @brief 低通滤波函数
    * @param self data cutoff fs order=5
    * @retval 无
    */
    '''
    def butter_lowpass_filter(self,data, cutoff, fs, order=5):
        b, a = self.butter_lowpass(cutoff, fs, order=order)
        y = scipy.signal.lfilter(b, a, data)
        return y  # Filter requirements.
    '''
    /*
    * @brief 低通滤波函数
    * @param self data cutoff fs order=5
    * @retval 无
    */
    '''
    def report_defect_docx_save(self,input_data):
        data_len = len(input_data)//4
        print(len(input_data))
        print(data_len)
        defect_number = 0
        i = 0
        while i < data_len:
            if input_data[4*i] > self.set_up_damage_limit or input_data[4*i] < self.set_down_damage_limit or input_data[4*i+1] > self.set_up_damage_limit or input_data[4*i+1] < self.set_down_damage_limit:
                defect_number += 1
                i += 800
            i += 1
        print(defect_number)
        doc = docx.Document()
        doc.add_heading('钢丝绳漏磁检测报告', 0)
        p = doc.add_paragraph('检测长度为0米')
        doc.add_paragraph('共有'+str(defect_number)+'处损伤')
        doc.add_paragraph('损伤位置为0米')
        doc.save("C:\\detec_data\\钢丝绳漏磁检测报表.docx")
        return defect_number
    '''
    /*
    * @brief 更新图像函数
    * @param self
    * @retval 无
    */
    '''
    def graph_update(self):
        if self.tcp_queue_user1_data.empty() == False:
            queue_get_data = self.tcp_queue_user1_data.get()
            self.graph_update_plot_func(queue_get_data)
            print("update")
        else:
            pass
'''
/*
* @brief 多线程类，用于socket通信数据获取
* @param PySide2.QtCore.QThread
* @retval 无
*/
'''
class tcp_commucate_qtThread(PySide2.QtCore.QThread):
    def __init__(self):
        super(tcp_commucate_qtThread,self).__init__()
    def run(self):
        main_win_obj.thread_tcp_commucate_get_func()
'''
/*
* @brief 多线程类，用于损伤判断
* @param PySide2.QtCore.QThread
* @retval 无
*/
'''
class defect_judge_qtThread(PySide2.QtCore.QThread):
    def __init__(self):
        super(defect_judge_qtThread,self).__init__()
    def run(self):
        main_win_obj.thread_defect_judge_func()
        #main_win_obj.thread_tcp_commucate_get_func()
'''
/*
* @brief matplotlib绘图类，用于测试
* @param FigureCanvasQTAgg
* @retval 无
*/
'''
class MplCanvas(FigureCanvasQTAgg):
    def __init__(self, parent=None, width=5, height=4, dpi=100,facecolor = '#052149'):
        self.fig = Figure(figsize=(width, height), dpi=dpi)
        super(MplCanvas, self).__init__(self.fig)


'''
/*
* @brief 主程序，程序由此执行
* @param 无
* @retval 无
*/
'''
if __name__ == '__main__':
    app = PySide2.QtWidgets.QApplication([])
    main_win_obj = main_application_obj()
    app.exec_()

